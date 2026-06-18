# -*- coding: utf-8 -*-
import os
import sys
import subprocess

def create_html_file(output_path):
    html_content = """<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>בדיקת כדאיות השקעה - מדריך ותרגילים</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Assistant:wght@300;400;600;700;800&display=swap');
        
        body {
            font-family: 'Assistant', sans-serif;
            background-color: #f8fafc;
            color: #1e293b;
            line-height: 1.6;
            margin: 0;
            padding: 40px 20px;
        }
        
        .container {
            max-width: 950px;
            margin: 0 auto;
            background: #ffffff;
            padding: 40px;
            border-radius: 16px;
            box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.05), 0 8px 10px -6px rgba(0, 0, 0, 0.05);
            border: 1px solid #e2e8f0;
        }
        
        h1 {
            color: #1e3a8a;
            font-size: 2.2rem;
            border-bottom: 3px solid #3b82f6;
            padding-bottom: 10px;
            margin-top: 0;
            font-weight: 800;
        }
        
        h2 {
            color: #2563eb;
            font-size: 1.6rem;
            margin-top: 30px;
            border-bottom: 1px solid #e2e8f0;
            padding-bottom: 5px;
            font-weight: 700;
        }
        
        h3 {
            color: #1d4ed8;
            font-size: 1.25rem;
            margin-top: 25px;
            font-weight: 600;
        }
        
        p {
            font-size: 1.1rem;
            margin-bottom: 15px;
        }
        
        ul, ol {
            margin-bottom: 20px;
            padding-right: 25px;
        }
        
        li {
            margin-bottom: 8px;
            font-size: 1.05rem;
        }
        
        .grid-container {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
            margin: 25px 0;
        }
        
        .card {
            padding: 20px;
            border-radius: 12px;
            border: 1px solid #e2e8f0;
        }
        
        .card-in {
            background-color: #f0fdf4;
            border-right: 5px solid #22c55e;
        }
        
        .card-in h3 {
            color: #15803d;
            margin-top: 0;
        }
        
        .card-out {
            background-color: #fef2f2;
            border-right: 5px solid #ef4444;
        }
        
        .card-out h3 {
            color: #b91c1c;
            margin-top: 0;
        }
        
        .alert {
            background-color: #eff6ff;
            border-right: 5px solid #3b82f6;
            padding: 15px;
            border-radius: 8px;
            margin: 20px 0;
        }
        
        .alert-title {
            font-weight: bold;
            color: #1e40af;
            margin-bottom: 5px;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 25px 0;
            font-size: 0.95rem;
            text-align: right;
        }
        
        th {
            background-color: #1e3a8a;
            color: #ffffff;
            font-weight: bold;
            padding: 12px;
            border: 1px solid #cbd5e1;
        }
        
        td {
            padding: 10px 12px;
            border: 1px solid #e2e8f0;
        }
        
        tr:nth-child(even) {
            background-color: #f8fafc;
        }
        
        .negative {
            color: #b91c1c;
            font-weight: 500;
            direction: ltr;
            text-align: right;
        }
        
        .positive {
            color: #15803d;
            font-weight: 500;
        }
        
        .formula {
            background: #f1f5f9;
            padding: 12px;
            border-radius: 6px;
            font-family: monospace;
            font-size: 1.05rem;
            display: inline-block;
            margin: 10px 0;
            direction: ltr;
        }
        
        .exercise-box {
            background: #ffffff;
            border: 2px solid #e2e8f0;
            border-radius: 12px;
            padding: 25px;
            margin: 30px 0;
        }
        
        .exercise-title {
            font-size: 1.3rem;
            font-weight: 700;
            color: #1e3a8a;
            margin-top: 0;
            border-bottom: 2px solid #3b82f6;
            padding-bottom: 8px;
            display: inline-block;
        }
        
        .solution-box {
            background: #fafafa;
            border: 1px dashed #cbd5e1;
            border-radius: 8px;
            padding: 20px;
            margin-top: 15px;
        }
        
        .solution-title {
            font-weight: bold;
            color: #475569;
            margin-bottom: 10px;
        }
        
        @media print {
            body {
                background-color: #ffffff;
                padding: 0;
            }
            .container {
                box-shadow: none;
                border: none;
                padding: 0;
            }
            .exercise-box {
                page-break-inside: avoid;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>בדיקת כדאיות השקעה: מדריך תיאורטי ותרגילים מעשיים</h1>
        <p>חומר עזר ותירגול לסטודנטים במימון. המדריך מפרט את הכללים לקביעת תזרימי מזומנים רלוונטיים ומציג 5 תרגילי כיתה פתורים בדרגות קושי שונות, כולל חישובי NPV.</p>
        
        <div class="alert">
            <div class="alert-title">עיקרון העל: עיקרון התזרים ההפרשי (Incremental Cash Flows)</div>
            אנו כוללים בניתוח אך ורק את תזרימי המזומנים שישתנו בחברה כתוצאה ישירה מקבלת הפרויקט (תזרימים שינבעו מהפרויקט בהשוואה למצב בלעדיו).
        </div>

        <h2>חלק 1: כללי אצבע - מה נכנס ומה לא נכנס לתחשיב כדאיות?</h2>
        
        <div class="grid-container">
            <!-- מה נכנס -->
            <div class="card card-in">
                <h3>✓ מה נכנס לתחשיב? (תזרימים רלוונטיים)</h3>
                <ul>
                    <li><strong>עלויות השקעה ראשוניות (CapEx):</strong> מחיר רכישת המכונה, הובלה, התקנה, הרצה והדרכת עובדים. כל אלו מהווים את בסיס הפחת של הנכס.</li>
                    <li><strong>שינוי בהון חוזר נקי (NWC):</strong> הגידול הנדרש במלאי או באשראי לקוחות (בניכוי אשראי ספקים) לצורך תמיכה בפרויקט. הון חוזר זה מוחזר במלואו כתזרים חיובי בשנה האחרונה.</li>
                    <li><strong>עלויות אלטרנטיביות (Opportunity Costs):</strong> שווי שוק של נכס קיים בבעלות החברה שישמש לפרויקט (לדוגמה, קרקע שניתן היה למכור או להשכיר). מחושב כתזרים שלילי נטו ממס בשנה 0.</li>
                    <li><strong>קניבליזציה והשפעות לוואי:</strong> ירידה צפויה במכירות של מוצרים קיימים של החברה עקב כניסת המוצר החדש.</li>
                    <li><strong>מגן מס מפחת:</strong> הפחת אינו תזרים מזומנים, אך הוא מקטין את חבות המס. מגן המס שווה ל-<em>פחת שנתי × שיעור המס</em>.</li>
                    <li><strong>תזרים גריטה נטו ממס (Salvage Value):</strong> המזומנים שיתקבלו ממכירת הנכס בסוף הפרויקט, בניכוי מס על רווח ההון (או בתוספת מגן מס במקרה של הפסד הון).</li>
                </ul>
            </div>
            
            <!-- מה לא נכנס -->
            <div class="card card-out">
                <h3>✗ מה לא נכנס לתחשיב? (תזרימים לא רלוונטיים)</h3>
                <ul>
                    <li><strong>עלויות שקועות (Sunk Costs):</strong> עלויות שכבר שולמו בעבר או שהחברה מחויבת להן בכל מקרה (למשל: סקר שוק, מחקר ופיתוח שכבר בוצע).</li>
                    <li><strong>הוצאות מימון (Financing Costs):</strong> תשלומי ריבית, החזרי קרן הלוואה, או חלוקת דיבידנדים.
                        <br><small style="color: #64748b;">(מדוע? כיוון שעלויות אלו משוקללות כבר בתוך שיעור ההיוון/מחיר ההון. הכללתן בתזרים תגרום לספירה כפולה).</small>
                    </li>
                    <li><strong>הקצאת עלויות תקורה קיימות:</strong> העמסת הוצאות קבועות של החברה (כגון שכירות משרדי הנהלה, משכורות מנכ"ל קיים) שאינן משתנות בעקבות הפרויקט.</li>
                    <li><strong>הוצאות פחת כשלעצמן:</strong> הפחת אינו יציאת מזומן פיזית. אנו משתמשים בו רק לצורך חישוב המס, ואז מוסיפים אותו בחזרה לרווח הנקי כדי לקבל את תזרים המזומנים התפעולי.</li>
                </ul>
            </div>
        </div>

        <h2>חלק 2: תרגילי כיתה פתורים</h2>

        <!-- תרגיל 1 -->
        <div class="exercise-box">
            <div class="exercise-title">תרגיל 1: פרויקט הרחבה (הרמה הבסיסית)</div>
            <p>חברת "סלטי השף" שוקלת לרכוש קו ייצור חדש בעלות של 500,000 ש"ח. הובלה והתקנה יעלו עוד 50,000 ש"ח.</p>
            <ul>
                <li><strong>אורך הפרויקט:</strong> 4 שנים.</li>
                <li><strong>פחת:</strong> המכונה מופחתת בשיטת הקו הישר לאורך 5 שנים לערך ספרים אפס (קצב פחת של 20% מהעלות המקורית בשנה).</li>
                <li><strong>הון חוזר:</strong> בשנה 0 נדרש להגדיל את המלאי ב-40,000 ש"ח. הון חוזר זה יוחזר במלואו בסוף שנה 4.</li>
                <li><strong>תוצאות תפעוליות:</strong> הכנסות שנתיות נוספות של 300,000 ש"ח, והוצאות תפעוליות שנתיות (ללא פחת) של 120,000 ש"ח.</li>
                <li>סקר שוק שנערך בחודש שעבר בעלות של 30,000 ש"ח הראה שיש ביקוש למוצר.</li>
                <li>בסוף שנה 4 המכונה תימכר לצד ג' תמורת 150,000 ש"ח.</li>
                <li>שיעור מס חברות ומס רווח הון: 25%.</li>
            </ul>
            
            <div class="solution-box">
                <div class="solution-title">פתרון תרגיל 1:</div>
                <ol>
                    <li><strong>בסיס הפחת של המכונה:</strong> 500,000 ש"ח (עלות המכונה) + 50,000 ש"ח (הובלה והתקנה) = 550,000 ש"ח.</li>
                    <li><strong>הפחת השנתי:</strong> 550,000 ש"ח × 20% = 110,000 ש"ח לשנה.</li>
                    <li><strong>עלות שקועה:</strong> סקר השוק בסך 30,000 ש"ח אינו רלוונטי ואינו נכלל בתזרים.</li>
                    <li><strong>ערך פנקסי (ספרים) בסוף שנה 4:</strong> 550,000 ש"ח (בסיס) - (110,000 ש"ח פחת × 4 שנים) = 110,000 ש"ח.</li>
                    <li><strong>מס על רווח הון מהגריטה:</strong>
                        <ul>
                            <li>רווח ההון: 150,000 ש"ח (מחיר מכירה) - 110,000 ש"ח (ערך ספרים) = 40,000 ש"ח.</li>
                            <li>המס: 40,000 ש"ח × 25% = 10,000 ש"ח.</li>
                            <li>תזרים גריטה נטו: 150,000 ש"ח - 10,000 ש"ח = 140,000 ש"ח.</li>
                        </ul>
                    </li>
                </ol>

                <table>
                    <thead>
                        <tr>
                            <th>סעיף (בשקלים)</th>
                            <th>שנה 0</th>
                            <th>שנה 1</th>
                            <th>שנה 2</th>
                            <th>שנה 3</th>
                            <th>שנה 4</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>השקעה במכונה והתקנה</td>
                            <td class="negative">-550,000</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                        </tr>
                        <tr>
                            <td>שינוי בהון חוזר נקי</td>
                            <td class="negative">-40,000</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td class="positive">40,000</td>
                        </tr>
                        <tr>
                            <td>הכנסות תפעוליות</td>
                            <td>-</td>
                            <td class="positive">300,000</td>
                            <td class="positive">300,000</td>
                            <td class="positive">300,000</td>
                            <td class="positive">300,000</td>
                        </tr>
                        <tr>
                            <td>הוצאות תפעוליות</td>
                            <td>-</td>
                            <td class="negative">-120,000</td>
                            <td class="negative">-120,000</td>
                            <td class="negative">-120,000</td>
                            <td class="negative">-120,000</td>
                        </tr>
                        <tr>
                            <td>פחת (לצרכי מס בלבד)</td>
                            <td>-</td>
                            <td style="color: #64748b;">-110,000</td>
                            <td style="color: #64748b;">-110,000</td>
                            <td style="color: #64748b;">-110,000</td>
                            <td style="color: #64748b;">-110,000</td>
                        </tr>
                        <tr>
                            <td>רווח תפעולי לפני מס (EBIT)</td>
                            <td>-</td>
                            <td>70,000</td>
                            <td>70,000</td>
                            <td>70,000</td>
                            <td>70,000</td>
                        </tr>
                        <tr>
                            <td>מס חברות (25%)</td>
                            <td>-</td>
                            <td class="negative">-17,500</td>
                            <td class="negative">-17,500</td>
                            <td class="negative">-17,500</td>
                            <td class="negative">-17,500</td>
                        </tr>
                        <tr>
                            <td>רווח נקי</td>
                            <td>-</td>
                            <td>52,500</td>
                            <td>52,500</td>
                            <td>52,500</td>
                            <td>52,500</td>
                        </tr>
                        <tr>
                            <td>הוספת פחת בחזרה</td>
                            <td>-</td>
                            <td class="positive">110,000</td>
                            <td class="positive">110,000</td>
                            <td class="positive">110,000</td>
                            <td class="positive">110,000</td>
                        </tr>
                        <tr style="background-color: #f1f5f9; font-weight: bold;">
                            <td>תזרים תפעולי (OCF)</td>
                            <td>-</td>
                            <td class="positive">162,500</td>
                            <td class="positive">162,500</td>
                            <td class="positive">162,500</td>
                            <td class="positive">162,500</td>
                        </tr>
                        <tr>
                            <td>תזרים גריטה נטו ממס</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td class="positive">140,000</td>
                        </tr>
                        <tr style="background-color: #cbd5e1; font-weight: bold;">
                            <td>תזרים מזומנים חופשי (FCF)</td>
                            <td class="negative">-590,000</td>
                            <td class="positive">162,500</td>
                            <td class="positive">162,500</td>
                            <td class="positive">162,500</td>
                            <td class="positive">342,500</td>
                        </tr>
                    </tbody>
                </table>
            </div>
        </div>

        <!-- תרגיל 2 -->
        <div class="exercise-box">
            <div class="exercise-title">תרגיל 2: פרויקט החלפת מכונה (הרמה הבינונית)</div>
            <p>חברה תעשייתית שוקלת להחליף מכונה ישנה במכונה חדשה ומתקדמת יותר.</p>
            <ul>
                <li><strong>המכונה הישנה:</strong> נרכשה לפני שנתיים ב-200,000 ש"ח. היא מופחתת בשיטת הקו הישר לאורך 5 שנים (40,000 ש"ח פחת לשנה). כיום ניתן למכור אותה תמורת 90,000 ש"ח. אם תישאר, ערך גריטה שלה בעוד 3 שנים יהיה אפס.</li>
                <li><strong>המכונה החדשה:</strong> מחיר רכישה 300,000 ש"ח. תופחת לאורך 3 שנים (אורך חיי הפרויקט) לערך ספרים אפס (100,000 ש"ח לשנה). בסוף שנה 3 היא תימכר ב-50,000 ש"ח.</li>
                <li><strong>שינוי תפעולי:</strong> המכונה החדשה תחסוך לחברה 80,000 ש"ח בשנה בהוצאות התפעול.</li>
                <li>שיעור מס חברות ומס רווח הון: 25%.</li>
            </ul>
            
            <div class="solution-box">
                <div class="solution-title">פתרון תרגיל 2:</div>
                <p>בפרויקט החלפה אנו בונים תזרים הפרשי: <strong>תזרים עם המכונה החדשה פחות תזרים ללא המכונה החדשה (המכונה הישנה).</strong></p>
                <ol>
                    <li><strong>השקעה נטו בשנה 0:</strong>
                        <ul>
                            <li>רכישת מכונה חדשה: 300,000- ש"ח.</li>
                            <li>מכירת המכונה הישנה: +90,000 ש"ח.</li>
                            <li>חישוב מס על מכירת הישנה: ערך פנקסי כיום הוא 120,000 ש"ח (עלות 200,000 פחות פחת נצבר לשנתיים של 80,000). מכיוון שהמכונה נמכרת ב-90,000 ש"ח, נוצר <strong>הפסד הון</strong> בסך 30,000 ש"ח. הפסד זה מקנה מגן מס (תזרים חיובי) של: 30,000 × 25% = 7,500+ ש"ח.</li>
                            <li>סך הכל השקעה נטו בשנה 0: -300,000 + 90,000 + 7,500 = <strong>202,500- ש"ח.</strong></li>
                        </ul>
                    </li>
                    <li><strong>חישוב פחת הפרשי (שנים 1-3):</strong>
                        <ul>
                            <li>פחת מכונה חדשה: 100,000 ש"ח לשנה.</li>
                            <li>פחת מכונה ישנה (שנחסך ונמנע מאיתנו): 40,000 ש"ח לשנה.</li>
                            <li>פחת הפרשי: 100,000 - 40,000 = 60,000 ש"ח לשנה.</li>
                        </ul>
                    </li>
                    <li><strong>תזרים תפעולי הפרשי (שנים 1-3):</strong>
                        <ul>
                            <li>חיסכון תפעולי שנתי: 80,000 ש"ח.</li>
                            <li>פחת הפרשי: 60,000- ש"ח.</li>
                            <li>גידול ברווח לצורכי מס: 20,000 ש"ח.</li>
                            <li>מס (25%): 5,000- ש"ח.</li>
                            <li>הוספת פחת הפרשי בחזרה: 60,000 ש"ח.</li>
                            <li>תזרים תפעולי הפרשי (OCF): 20,000 - 5,000 + 60,000 = <strong>75,000 ש"ח.</strong></li>
                        </ul>
                    </li>
                    <li><strong>תזרים גריטה בשנה 3:</strong>
                        <ul>
                            <li>מכירת המכונה החדשה: 50,000 ש"ח.</li>
                            <li>ערך ספרים: 0 ש"ח.</li>
                            <li>מס רווח הון: 50,000 × 25% = 12,500- ש"ח.</li>
                            <li>גריטה נטו: 50,000 - 12,500 = <strong>37,500 ש"ח.</strong></li>
                        </ul>
                    </li>
                </ol>

                <table>
                    <thead>
                        <tr>
                            <th>תזרים הפרשי (בשקלים)</th>
                            <th>שנה 0</th>
                            <th>שנה 1</th>
                            <th>שנה 2</th>
                            <th>שנה 3</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>רכישת מכונה חדשה</td>
                            <td class="negative">-300,000</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                        </tr>
                        <tr>
                            <td>מכירת מכונה ישנה (כולל מגן מס)</td>
                            <td class="positive">97,500</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                        </tr>
                        <tr>
                            <td>חיסכון תפעולי שנתי בהוצאות</td>
                            <td>-</td>
                            <td class="positive">80,000</td>
                            <td class="positive">80,000</td>
                            <td class="positive">80,000</td>
                        </tr>
                        <tr>
                            <td>מס חברות על החיסכון והפחת ההפרשי</td>
                            <td>-</td>
                            <td class="negative">-5,000</td>
                            <td class="negative">-5,000</td>
                            <td class="negative">-5,000</td>
                        </tr>
                        <tr style="background-color: #f1f5f9; font-weight: bold;">
                            <td>תזרים תפעולי הפרשי נטו (OCF)</td>
                            <td>-</td>
                            <td class="positive">75,000</td>
                            <td class="positive">75,000</td>
                            <td class="positive">75,000</td>
                        </tr>
                        <tr>
                            <td>גריטה נטו של המכונה החדשה</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td class="positive">37,500</td>
                        </tr>
                        <tr style="background-color: #cbd5e1; font-weight: bold;">
                            <td>תזרים מזומנים חופשי הפרשי (FCF)</td>
                            <td class="negative">-202,500</td>
                            <td class="positive">75,000</td>
                            <td class="positive">75,000</td>
                            <td class="positive">112,500</td>
                        </tr>
                    </tbody>
                </table>
            </div>
        </div>

        <!-- תרגיל 3 -->
        <div class="exercise-box">
            <div class="exercise-title">תרגיל 3: עלות אלטרנטיבית וקניבליזציה (הרמה המתקדמת)</div>
            <p>חברת "טק-פליי" רוצה להשיק קונסולת משחקים חדשה, בשם "פליי-קיד".</p>
            <ul>
                <li><strong>השקעה במפעל וציוד:</strong> 800,000 ש"ח בשנה 0. מופחת בשיטת הקו הישר ל-4 שנים לערך ספרים אפס (200,000 ש"ח לשנה).</li>
                <li><strong>עלות אלטרנטיבית:</strong> הפרויקט ימוקם במבנה קיים של החברה. החברה יכלה להשכיר מבנה זה לחברה חיצונית תמורת 50,000 ש"ח בשנה (משולם בתחילת כל שנה, כלומר בשנים 0, 1, 2 ו-3).</li>
                <li><strong>קניבליזציה:</strong> השקת הקונסולת החדשה תפגע במכירות של קונסולה ישנה של החברה ותקטין את תזרים המזומנים התפעולי של החברה ב-30,000 ש"ח בשנה (בשנים 1 עד 4).</li>
                <li><strong>תזרים תפעולי מהקונסולה החדשה:</strong> הכנסות שנתיות של 400,000 ש"ח והוצאות תפעוליות שנתיות (ללא פחת) של 150,000 ש"ח.</li>
                <li>שיעור מס חברות: 20% (חל גם על הכנסות השכירות).</li>
            </ul>
            
            <div class="solution-box">
                <div class="solution-title">פתרון תרגיל 3:</div>
                <ol>
                    <li><strong>עלות אלטרנטיבית (שכירות נטו ממס):</strong> 
                        הפסד השכירות ברוטו הוא 50,000 ש"ח. היות והחברה הייתה משלמת 20% מס על הכנסה זו, הפסד המזומנים נטו הוא: $50,000 \times (1 - 0.20) = 40,000$ ש"ח בשנה.
                        עלות זו תירשם כתזרים שלילי בשנים 0 עד 3 (במועד קבלת השכירות).
                    </li>
                    <li><strong>קניבליזציה:</strong>
                        הנתון מציין אובדן תזרים מזומנים תפעולי בסך 30,000 ש"ח (לאחר השפעת המס). נרשום זאת כתזרים שלילי בשנים 1 עד 4.
                    </li>
                    <li><strong>חישוב תזרים תפעולי מהקונסולה החדשה (OCF):</strong>
                        <ul>
                            <li>EBITDA לשנה: $400,000 - 150,000 = 250,000$ ש"ח.</li>
                            <li>פחת שנתי: $200,000$ ש"ח.</li>
                            <li>רווח לפני מס: $250,000 - 200,000 = 50,000$ ש"ח.</li>
                            <li>מס (20%): $10,000$ ש"ח.</li>
                            <li>רווח נקי: $40,000$ ש"ח.</li>
                            <li>הוספת פחת חזרה: $200,000$ ש"ח.</li>
                            <li>OCF שנתי מהקונסולה: $40,000 + 200,000 = 240,000$ ש"ח.</li>
                        </ul>
                    </li>
                </ol>

                <table>
                    <thead>
                        <tr>
                            <th>סעיף (בשקלים)</th>
                            <th>שנה 0</th>
                            <th>שנה 1</th>
                            <th>שנה 2</th>
                            <th>שנה 3</th>
                            <th>שנה 4</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>השקעה בציוד ומפעל</td>
                            <td class="negative">-800,000</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                        </tr>
                        <tr>
                            <td>עלות אלטרנטיבית (שכירות נטו)</td>
                            <td class="negative">-40,000</td>
                            <td class="negative">-40,000</td>
                            <td class="negative">-40,000</td>
                            <td class="negative">-40,000</td>
                            <td>-</td>
                        </tr>
                        <tr>
                            <td>תזרים תפעולי מקונסולה חדשה</td>
                            <td>-</td>
                            <td class="positive">240,000</td>
                            <td class="positive">240,000</td>
                            <td class="positive">240,000</td>
                            <td class="positive">240,000</td>
                        </tr>
                        <tr>
                            <td>קניבליזציה (אובדן תזרים ישן)</td>
                            <td>-</td>
                            <td class="negative">-30,000</td>
                            <td class="negative">-30,000</td>
                            <td class="negative">-30,000</td>
                            <td class="negative">-30,000</td>
                        </tr>
                        <tr style="background-color: #cbd5e1; font-weight: bold;">
                            <td>תזרים מזומנים חופשי (FCF)</td>
                            <td class="negative">-840,000</td>
                            <td class="positive">170,000</td>
                            <td class="positive">170,000</td>
                            <td class="positive">170,000</td>
                            <td class="positive">210,000</td>
                        </tr>
                    </tbody>
                </table>
            </div>
        </div>

        <!-- תרגיל 4 -->
        <div class="exercise-box">
            <div class="exercise-title">תרגיל 4: השקת מוצר חדש עם חישוב NPV (חדש!)</div>
            <p>חברת "גיימינג פרו" שוקלת לרכוש ציוד לייצור בקרים חדשים. עלות הציוד היא 1,200,000 ש"ח, ועלויות התקנה והרצה הן 100,000 ש"ח.</p>
            <ul>
                <li><strong>אורך הפרויקט:</strong> 5 שנים.</li>
                <li><strong>פחת:</strong> שיטת קו ישר לאורך 5 שנים לערך ספרים 0.</li>
                <li><strong>עלות שקועה:</strong> סקר שוק בעלות 50,000 ש"ח שבוצע בחודש שעבר.</li>
                <li><strong>הון חוזר נקי (NWC):</strong> נדרשת השקעה ראשונית של 80,000 ש"ח בשנה 0, אשר תוחזר במלואה בסוף שנה 5.</li>
                <li><strong>נתונים תפעוליים:</strong> הכנסות של 600,000 ש"ח בשנה, והוצאות תפעוליות (ללא פחת) של 200,000 ש"ח בשנה.</li>
                <li><strong>קניבליזציה:</strong> ההשקה תפגע במוצרים אחרים ותפחית את תזרים המזומנים התפעולי של החברה ב-20,000 ש"ח בשנה (לאחר מס).</li>
                <li><strong>גריטה:</strong> הציוד יימכר בסוף שנה 5 ב-200,000 ש"ח.</li>
                <li><strong>שיעור המס:</strong> 23%.</li>
                <li><strong>מחיר ההון של הפרויקט (שיעור ההיוון):</strong> 12%.</li>
            </ul>
            
            <div class="solution-box">
                <div class="solution-title">פתרון תרגיל 4:</div>
                <ol>
                    <li><strong>שנה 0 (השקעה ראשונית נטו):</strong> 
                        $-1,200,000 \text{ (ציוד)} - 100,000 \text{ (התקנה)} - 80,000 \text{ (הון חוזר)} = \mathbf{-1,380,000 \text{ ש\"ח}}$. 
                        סקר השוק (50,000 ש"ח) הוא עלות שקועה ואינו נכלל.
                    </li>
                    <li><strong>פחת שנתי:</strong> $1,300,000 / 5 = 260,000$ ש"ח לשנה.</li>
                    <li><strong>תזרים תפעולי (שנים 1-5):</strong>
                        <ul>
                            <li>EBITDA מהמוצר החדש: $600,000 - 200,000 = 400,000$ ש"ח.</li>
                            <li>רווח לפני מס (EBIT): $400,000 - 260,000 \text{ (פחת)} = 140,000$ ש"ח.</li>
                            <li>מס (23%): $-32,200$ ש"ח.</li>
                            <li>רווח נקי: $107,800$ ש"ח.</li>
                            <li>הוספת פחת: $+260,000$ ש"ח.</li>
                            <li>OCF מהבקרים: $367,800$ ש"ח.</li>
                            <li>בניכוי קניבליזציה: $367,800 - 20,000 = \mathbf{347,800 \text{ ש\"ח}}$.</li>
                        </ul>
                    </li>
                    <li><strong>תזרים גריטה נטו בשנה 5:</strong>
                        ערך ספרים הוא 0 ש"ח. רווח הון = 200,000 ש"ח. מס = $200,000 \times 23\% = 46,000$ ש"ח.
                        תזרים גריטה נטו: $200,000 - 46,000 = \mathbf{154,000 \text{ ש\"ח}}$.
                    </li>
                    <li><strong>שנה 5 FCF כולל:</strong> $347,800 \text{ (תפעולי)} + 154,000 \text{ (גריטה)} + 80,000 \text{ (החזר הון חוזר)} = \mathbf{581,800 \text{ ש\"ח}}$.</li>
                    <li><strong>חישוב NPV ב-12%:</strong>
                        $$\text{PV} = \frac{347,800}{1.12^1} + \frac{347,800}{1.12^2} + \frac{347,800}{1.12^3} + \frac{347,800}{1.12^4} + \frac{581,800}{1.12^5}$$
                        $$\text{PV} = 310,536 + 277,264 + 247,557 + 221,033 + 330,129 = \mathbf{1,386,519 \text{ ש\"ח}}$$
                        $$\text{NPV} = 1,386,519 - 1,380,000 \text{ (השקעה שנת 0)} = \mathbf{+6,519 \text{ ש\"ח}}$$
                        <strong>החלטה:</strong> ה-NPV חיובי ולכן הפרויקט כדאי ומגדיל את ערך החברה ב-6,519 ש"ח.
                    </li>
                </ol>

                <table>
                    <thead>
                        <tr>
                            <th>סעיף (בשקלים)</th>
                            <th>שנה 0</th>
                            <th>שנה 1</th>
                            <th>שנה 2</th>
                            <th>שנה 3</th>
                            <th>שנה 4</th>
                            <th>שנה 5</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>ציוד והתקנה</td>
                            <td class="negative">-1,300,000</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                        </tr>
                        <tr>
                            <td>השקעה בהון חוזר נקי</td>
                            <td class="negative">-80,000</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td class="positive">80,000</td>
                        </tr>
                        <tr>
                            <td>תזרים תפעולי (OCF) נטו</td>
                            <td>-</td>
                            <td class="positive">347,800</td>
                            <td class="positive">347,800</td>
                            <td class="positive">347,800</td>
                            <td class="positive">347,800</td>
                            <td class="positive">347,800</td>
                        </tr>
                        <tr>
                            <td>גריטה נטו ממס</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td class="positive">154,000</td>
                        </tr>
                        <tr style="background-color: #cbd5e1; font-weight: bold;">
                            <td>תזרים FCF חופשי</td>
                            <td class="negative">-1,380,000</td>
                            <td class="positive">347,800</td>
                            <td class="positive">347,800</td>
                            <td class="positive">347,800</td>
                            <td class="positive">347,800</td>
                            <td class="positive">581,800</td>
                        </tr>
                    </tbody>
                </table>
            </div>
        </div>

        <!-- תרגיל 5 -->
        <div class="exercise-box">
            <div class="exercise-title">תרגיל 5: פרויקט החלפת ציוד עם חישוב NPV (חדש!)</div>
            <p>חברה שוקלת להחליף מכונה ישנה במכונה חדשה וחסכונית.</p>
            <ul>
                <li><strong>המכונה הישנה:</strong> נרכשה לפני 3 שנים ב-600,000 ש"ח. היא מופחתת בקו ישר ל-6 שנים (100,000 ש"ח לשנה). כיום ניתנת למכירה ב-250,000 ש"ח. אם תישאר, ערך הגריטה שלה בעוד 3 שנים יהיה 50,000 ש"ח.</li>
                <li><strong>המכונה החדשה:</strong> עלות רכישה והתקנה היא 850,000 ש"ח בשנה 0. מופחתת בקו ישר לאורך 3 שנים לערך ספרים של 50,000 ש"ח. בסוף שנה 3 היא תימכר ב-120,000 ש"ח.</li>
                <li><strong>שינוי תפעולי:</strong> המכונה החדשה תחסוך להוצאות התפעול 180,000 ש"ח בשנה (לפני מס).</li>
                <li><strong>שיעור מס:</strong> 20%.</li>
                <li><strong>מחיר ההון (שיעור ההיוון):</strong> 10%.</li>
            </ul>
            
            <div class="solution-box">
                <div class="solution-title">פתרון תרגיל 5:</div>
                <ol>
                    <li><strong>שנה 0 (השקעה נטו להחלפה):</strong> 
                        רכישת חדשה: $-850,000$ ש"ח. מכירת ישנה: $+250,000$ ש"ח.
                        ערך ספרים ישנה: $600,000 - (3 \times 100,000) = 300,000$ ש"ח. 
                        מכירת הישנה ב-250,000 ש"ח יוצרת <strong>הפסד הון</strong> של 50,000 ש"ח שמקנה מגן מס (תזרים חיובי): $50,000 \times 20\% = +10,000$ ש"ח.
                        תזרים שנה 0: $-850,000 + 250,000 + 10,000 = \mathbf{-590,000 \text{ ש\"ח}}$.
                    </li>
                    <li><strong>פחת הפרשי (שנים 1-3):</strong>
                        פחת חדשה: $(850,000 - 50,000) / 3 = 266,667$ ש"ח לשנה. 
                        פחת ישנה (נחסך): $100,000$ ש"ח לשנה.
                        פחת הפרשי: $266,667 - 100,000 = 166,667$ ש"ח לשנה.
                    </li>
                    <li><strong>תזרים תפעולי הפרשי (OCF):</strong>
                        $$\text{OCF} = (\text{חיסכון} - \text{פחת הפרשי}) \times (1 - T) + \text{פחת הפרשי}$$
                        $$\text{OCF} = (180,000 - 166,667) \times 0.8 + 166,667 = \mathbf{177,333 \text{ ש\"ח}} \text{ לשנה}.$$
                    </li>
                    <li><strong>תזרים גריטה הפרשי נטו בשנה 3:</strong>
                        <ul>
                            <li>גריטה נטו מהמכונה החדשה: מכירה ב-120,000 ש"ח. ערך ספרים 50,000 ש"ח. רווח הון: 70,000 ש"ח. מס: $70,000 \times 20\% = 14,000$ ש"ח. תזרים נטו: $120,000 - 14,000 = 106,000$ ש"ח.</li>
                            <li>גריטה נטו שנמנעה מהמכונה הישנה: מכירה ב-50,000 ש"ח. ערך ספרים 0 ש"ח (עברו 6 שנים). רווח הון: 50,000 ש"ח. מס: $50,000 \times 20\% = 10,000$ ש"ח. תזרים נטו: $50,000 - 10,000 = 40,000$ ש"ח.</li>
                            <li>גריטה הפרשית נטו בשנה 3: $106,000 - 40,000 = \mathbf{66,000 \text{ ש\"ח}}$.</li>
                        </ul>
                    </li>
                    <li><strong>חישוב NPV הפרשי ב-10%:</strong>
                        $$\text{PV} = \frac{177,333}{1.1^1} + \frac{177,333}{1.1^2} + \frac{177,333 + 66,000 \text{ (גריטה הפרשית)}}{1.1^3}$$
                        $$\text{PV} = 161,212 + 146,556 + \frac{243,333}{1.1^3} = 161,212 + 146,556 + 182,820 = \mathbf{490,588 \text{ ש\"ח}}$$
                        $$\text{NPV} = 490,588 - 590,000 = \mathbf{-99,412 \text{ ש\"ח}}$$
                        <strong>החלטה:</strong> ה-NPV ההפרשי שלילי. ההחלפה אינה כדאית לחברה ומומלץ להמשיך להשתמש במכונה הישנה.
                    </li>
                </ol>

                <table>
                    <thead>
                        <tr>
                            <th>תזרים הפרשי (בשקלים)</th>
                            <th>שנה 0</th>
                            <th>שנה 1</th>
                            <th>שנה 2</th>
                            <th>שנה 3</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td>השקעה נטו במכונות</td>
                            <td class="negative">-590,000</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                        </tr>
                        <tr>
                            <td>תזרים תפעולי הפרשי (OCF)</td>
                            <td>-</td>
                            <td class="positive">177,333</td>
                            <td class="positive">177,333</td>
                            <td class="positive">177,333</td>
                        </tr>
                        <tr>
                            <td>גריטה הפרשית נטו</td>
                            <td>-</td>
                            <td>-</td>
                            <td>-</td>
                            <td class="positive">66,000</td>
                        </tr>
                        <tr style="background-color: #cbd5e1; font-weight: bold;">
                            <td>תזרים FCF חופשי הפרשי</td>
                            <td class="negative">-590,000</td>
                            <td class="positive">177,333</td>
                            <td class="positive">177,333</td>
                            <td class="positive">243,333</td>
                        </tr>
                    </tbody>
                </table>
            </div>
        </div>
    </div>
</body>
</html>
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"Created HTML file at: {output_path}")

def create_markdown_file(output_path):
    md_content = """# בדיקת כדאיות השקעה: מדריך תיאורטי ותרגילים מעשיים

חומר עזר ותירגול לסטודנטים במימון. המדריך מפרט את הכללים לקביעת תזרימי מזומנים רלוונטיים ומציג 5 תרגילי כיתה פתורים בדרגות קושי שונות, כולל חישובי NPV.

## עיקרון העל: עיקרון התזרים ההפרשי (Incremental Cash Flows)
אנו כוללים בניתוח אך ורק את תזרימי המזומנים שישתנו בחברה כתוצאה ישירה מקבלת הפרויקט (תזרימים שינבעו מהפרויקט בהשוואה למצב בלעדיו).

---

## חלק 1: כללי אצבע - מה נכנס ומה לא נכנס לתחשיב כדאיות?

### ✓ מה נכנס לתחשיב? (תזרימים רלוונטיים)
1. **עלויות השקעה ראשוניות (CapEx):** מחיר רכישת המכונה, הובלה, התקנה, הרצה והדרכת עובדים. כל אלו מהווים את בסיס הפחת של הנכס.
2. **שינוי בהון חוזר נקי (NWC):** הגידול הנדרש במלאי או באשראי לקוחות (בניכוי אשראי ספקים) לצורך תמיכה בפרויקט. הון חוזר זה מוחזר במלואו כתזרים חיובי בשנה האחרונה.
3. **עלויות אלטרנטיביות (Opportunity Costs):** שווי שוק של נכס קיים בבעלות החברה שישמש לפרויקט (לדוגמה, קרקע שניתן היה למכור או להשכיר). מחושב כתזרים שלילי נטו ממס בשנה 0.
4. **קניבליזציה והשפעות לוואי:** ירידה צפויה במכירות של מוצרים קיימים של החברה עקב כניסת המוצר החדש.
5. **מגן מס מפחת:** הפחת אינו תזרים מזומנים, אך הוא מקטין את חבות המס. מגן המס שווה ל- *פחת שנתי × שיעור המס*.
6. **תזרים גריטה נטו ממס (Salvage Value):** המזומנים שיתקבלו ממכירת הנכס בסוף הפרויקט, בניכוי מס על רווח ההון (או בתוספת מגן מס במקרה של הפסד הון).

### ✗ מה לא נכנס לתחשיב? (תזרימים לא רלוונטיים)
1. **עלויות שקועות (Sunk Costs):** עלויות שכבר שולמו בעבר או שהחברה מחויבת להן בכל מקרה (למשל: סקר שוק, מחקר ופיתוח שכבר בוצע).
2. **הוצאות מימון (Financing Costs):** תשלומי ריבית, החזרי קרן הלוואה, או חלוקת דיבידנדים.
   *(הסבר: עלויות אלו משוקללות כבר בתוך שיעור ההיוון/מחיר ההון (WACC). הכללתן בתזרים תגרום לספירה כפולה).*
3. **הקצאת עלויות תקורה קיימות:** העמסת הוצאות קבועות של החברה (כגון שכירות משרדי הנהלה, משכורות מנכ"ל קיים) שאינן משתנות בעקבות הפרויקט.
4. **הוצאות פחת כשלעצמן:** הפחת אינו יציאת מזומן פיזית. אנו משתמשים בו רק לצורך חישוב המס, ואז מוסיפים אותו בחזרה לרווח הנקי כדי לקבל את תזרים המזומנים התפעולי.

---

## חלק 2: תרגילי כיתה פתורים

### תרגיל 1: פרויקט הרחבה (הרמה הבסיסית)
חברת "סלטי השף" שוקלת לרכוש קו ייצור חדש בעלות של 500,000 ש"ח. הובלה והתקנה יעלו עוד 50,000 ש"ח.
* **אורך הפרויקט:** 4 שנים.
* **פחת:** המכונה מופחתת בשיטת הקו הישר לאורך 5 שנים לערך ספרים אפס.
* **הון חוזר:** בשנה 0 נדרש להגדיל את המלאי ב-40,000 ש"ח. הון חוזר זה יוחזר במלואו בסוף שנה 4.
* **תוצאות תפעוליות:** הכנסות שנתיות נוספות של 300,000 ש"ח, והוצאות תפעוליות שנתיות (ללא פחת) של 120,000 ש"ח.
* סקר שוק שנערך בחודש שעבר בעלות של 30,000 ש"ח הראה שיש ביקוש למוצר.
* בסוף שנה 4 המכונה תימכר לצד ג' תמורת 150,000 ש"ח.
* שיעור מס חברות ומס רווח הון: 25%.

#### פתרון תרגיל 1:
1. **בסיס הפחת:** 500,000 + 50,000 = 550,000 ש"ח.
2. **הפחת השנתי:** 550,000 × 20% = 110,000 ש"ח לשנה.
3. **עלות ספרים בסוף שנה 4:** 550,000 - (110,000 × 4) = 110,000 ש"ח.
4. **תזרים גריטה נטו:** 150,000 - 25% × (150,000 - 110,000) = 140,000 ש"ח.

| סעיף (בשקלים) | שנה 0 | שנה 1 | שנה 2 | שנה 3 | שנה 4 |
| :--- | :---: | :---: | :---: | :---: | :---: |
| השקעה במכונה והתקנה | -550,000 | - | - | - | - |
| שינוי בהון חוזר נקי | -40,000 | - | - | - | 40,000 |
| הכנסות תפעוליות | - | 300,000 | 300,000 | 300,000 | 300,000 |
| הוצאות תפעוליות | - | -120,000 | -120,000 | -120,000 | -120,000 |
| פחת (לצרכי מס בלבד) | - | -110,000 | -110,000 | -110,000 | -110,000 |
| רווח לפני מס (EBIT) | - | 70,000 | 70,000 | 70,000 | 70,000 |
| מס חברות (25%) | - | -17,500 | -17,500 | -17,500 | -17,500 |
| רווח נקי | - | 52,500 | 52,500 | 52,500 | 52,500 |
| הוספת פחת בחזרה | - | 110,000 | 110,000 | 110,000 | 110,000 |
| **תזרים תפעולי (OCF)** | - | **162,500** | **162,500** | **162,500** | **162,500** |
| תזרים גריטה נטו ממס | - | - | - | - | 140,000 |
| **תזרים מזומנים חופשי (FCF)** | **-590,000** | **162,500** | **162,500** | **162,500** | **342,500** |

---

### תרגיל 2: פרויקט החלפת מכונה (הרמה הבינונית)
* **המכונה הישנה:** נרכשה לפני שנתיים ב-200,000 ש"ח. מופחתת בקו ישר לאורך 5 שנים (40,000 ש"ח לשנה). כיום ניתנת למכירה ב-90,000 ש"ח.
* **המכונה החדשה:** מחיר רכישה 300,000 ש"ח. תופחת לאורך 3 שנים לערך ספרים אפס. בסוף שנה 3 תימכר ב-50,000 ש"ח.
* **שינוי תפעולי:** חיסכון בהוצאות תפעול של 80,000 ש"ח בשנה.
* שיעור מס: 25%.

#### פתרון תרגיל 2:
1. **השקעה נטו בשנה 0:** -300,000 + 90,000 + 25% × (120,000 - 90,000) = **202,500- ש"ח**.
2. **פחת הפרשי:** 100,000 - 40,000 = 60,000 ש"ח לשנה.
3. **תזרים תפעולי הפרשי (OCF):** (80,000 - 60,000) × 75% + 60,000 = **75,000 ש"ח**.
4. **גריטה נטו בשנה 3:** 50,000 - 25% × (50,000 - 0) = **37,500 ש"ח**.

| תזרים הפרשי (בשקלים) | שנה 0 | שנה 1 | שנה 2 | שנה 3 |
| :--- | :---: | :---: | :---: | :---: |
| השקעה נטו במכונות | -202,500 | - | - | - |
| תזרים תפעולי הפרשי (OCF) | - | 75,000 | 75,000 | 75,000 |
| גריטה נטו של המכונה החדשה | - | - | - | 37,500 |
| **תזרים מזומנים חופשי הפרשי (FCF)** | **-202,500** | **75,000** | **75,000** | **112,500** |

---

### תרגיל 3: עלות אלטרנטיבית וקניבליזציה (הרמה המתקדמת)
* **השקעה ראשונית:** 800,000 ש"ח בשנה 0. מופחת בקו ישר ל-4 שנים לערך ספרים אפס.
* **עלות אלטרנטיבית:** אובדן שכירות של 50,000 ש"ח בשנה (לפני מס) בשנים 0, 1, 2 ו-3.
* **קניבליזציה:** ירידה בתזרים המזומנים של מוצרים קיימים בסך 30,000 ש"ח בשנה (שנים 1-4).
* **תוצאות תפעוליות של החדש:** הכנסות 400,000 ש"ח, הוצאות תפעוליות 150,000 ש"ח.
* שיעור מס: 20%.

#### פתרון תרגיל 3:
1. **עלות אלטרנטיבית נטו:** 50,000 × (1 - 0.20) = **40,000- ש"ח** (שנים 0-3).
2. **קניבליזציה נטו:** **30,000- ש"ח** (שנים 1-4).
3. **OCF חדש:** (250,000 - 200,000) × 80% + 200,000 = **240,000 ש"ח**.

| סעיף (בשקלים) | שנה 0 | שנה 1 | שנה 2 | שנה 3 | שנה 4 |
| :--- | :---: | :---: | :---: | :---: | :---: |
| השקעה בציוד ומפעל | -800,000 | - | - | - | - |
| עלות אלטרנטיבית (שכירות נטו) | -40,000 | -40,000 | -40,000 | -40,000 | - |
| תזרים תפעולי מקונסולה חדשה | - | 240,000 | 240,000 | 240,000 | 240,000 |
| קניבליזציה | - | -30,000 | -30,000 | -30,000 | -30,000 |
| **תזרים מזומנים חופשי (FCF)** | **-840,000** | **170,000** | **170,000** | **170,000** | **210,000** |

---

### תרגיל 4: השקת מוצר חדש עם חישוב NPV (חדש!)
* **השקעה בציוד והתקנה:** 1,200,000 ש"ח רכישה + 100,000 ש"ח התקנה בשנה 0.
* **אורך הפרויקט:** 5 שנים.
* **פחת:** שיטת קו ישר לאורך 5 שנים לערך ספרים 0 (פחת שנתי של 260,000 ש"ח).
* **הון חוזר נקי (NWC):** נדרשת השקעה ראשונית של 80,000 ש"ח בשנה 0, המוחזרת במלואה בשנה 5.
* **נתונים תפעוליים:** הכנסות 600,000 ש"ח בשנה, והוצאות תפעוליות (ללא פחת) של 200,000 ש"ח בשנה.
* **קניבליזציה:** הפסד תזרים שנתי של מוצרים קיימים בגובה 20,000 ש"ח בשנה (לאחר מס).
* **גריטה:** הציוד יימכר בשנה 5 ב-200,000 ש"ח.
* **שיעור המס:** 23%.
* **מחיר ההון (שיעור ההיוון):** 12%.

#### פתרון תרגיל 4:
1. **השקעה ראשונית (שנה 0):** -1,200,000 (ציוד) - 100,000 (התקנה) - 80,000 (הון חוזר) = **1,380,000- ש"ח**.
2. **פחת שנתי:** 1,300,000 / 5 = 260,000 ש"ח.
3. **תזרים תפעולי (OCF) שנים 1-5:**
   * EBITDA: 400,000 ש"ח.
   * EBIT: 400,000 - 260,000 = 140,000 ש"ח.
   * מס: 140,000 × 23% = 32,200 ש"ח.
   * רווח נקי: 107,800 ש"ח.
   * OCF ללא קניבליזציה: 107,800 + 260,000 = 367,800 ש"ח.
   * OCF כולל קניבליזציה: 367,800 - 20,000 = **347,800 ש"ח**.
4. **גריטה נטו בשנה 5:** 200,000 - 23% × (200,000 - 0) = **154,000 ש"ח**.
5. **תזרים מזומנים בשנה 5 (FCF):** 347,800 (תפעולי) + 154,000 (גריטה) + 80,000 (החזר הון חוזר) = **581,800 ש"ח**.
6. **חישוב NPV ב-12%:**
   * $PV = \frac{347,800}{1.12^1} + \frac{347,800}{1.12^2} + \frac{347,800}{1.12^3} + \frac{347,800}{1.12^4} + \frac{581,800}{1.12^5}$
   * $PV = 310,536 + 277,264 + 247,557 + 221,033 + 330,129 = 1,386,519$ ש"ח.
   * $NPV = 1,386,519 - 1,380,000 = \mathbf{+6,519 \text{ ש\"ח}}$.
   * **החלטה:** לקבל את הפרויקט (NPV חיובי).

| סעיף (בשקלים) | שנה 0 | שנה 1 | שנה 2 | שנה 3 | שנה 4 | שנה 5 |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: |
| ציוד והתקנה | -1,300,000 | - | - | - | - | - |
| השקעה בהון חוזר נקי | -80,000 | - | - | - | - | 80,000 |
| תזרים תפעולי (OCF) נטו | - | 347,800 | 347,800 | 347,800 | 347,800 | 347,800 |
| גריטה נטו ממס | - | - | - | - | - | 154,000 |
| **תזרים FCF חופשי** | **-1,380,000** | **347,800** | **347,800** | **347,800** | **347,800** | **581,800** |

---

### תרגיל 5: פרויקט החלפת ציוד עם חישוב NPV (חדש!)
* **המכונה הישנה:** נרכשה לפני 3 שנים ב-600,000 ש"ח. מופחתת בקו ישר ל-6 שנים (100,000 ש"ח לשנה). כיום ניתנת למכירה ב-250,000 ש"ח. אם תישאר, גריטתה בעוד 3 שנים תהיה 50,000 ש"ח.
* **המכונה החדשה:** עלות רכישה והתקנה 850,000 ש"ח בשנה 0. מופחתת בקו ישר ל-3 שנים לערך ספרים של 50,000 ש"ח. בסוף שנה 3 תימכר ב-120,000 ש"ח.
* **שינוי תפעולי:** חיסכון בהוצאות תפעול של 180,000 ש"ח בשנה (לפני מס).
* **שיעור מס:** 20%.
* **מחיר ההון (שיעור ההיוון):** 10%.

#### פתרון תרגיל 5:
1. **שנה 0 (השקעה נטו להחלפה):**
   * רכישת חדשה: -850,000 ש"ח.
   * מכירת ישנה: +250,000 ש"ח.
   * מס על מכירת ישנה (הפסד הון): ערך ספרים של הישנה הוא 300,000 ש"ח. מכיוון שהיא נמכרת ב-250,000 ש"ח, נוצר הפסד הון של 50,000 ש"ח המניב מגן מס של 10,000 ש"ח.
   * תזרים שנה 0: -850,000 + 250,000 + 10,000 = **590,000- ש"ח**.
2. **פחת הפרשי:**
   * פחת חדשה: (850,000 - 50,000) / 3 = 266,667 ש"ח לשנה.
   * פחת ישנה (נחסך): 100,000 ש"ח לשנה.
   * פחת הפרשי: 266,667 - 100,000 = 166,667 ש"ח לשנה.
3. **תזרים תפעולי הפרשי (OCF):**
   * $(180,000 - 166,667) \times 80\% + 166,667 = \mathbf{177,333 \text{ ש\"ח}}$ לשנה.
4. **תזרים גריטה הפרשי נטו בשנה 3:**
   * גריטה נטו חדשה: 120,000 - 20% × (120,000 - 50,000) = 106,000 ש"ח.
   * גריטה נטו שנמנעה מהישנה: 50,000 - 20% × (50,000 - 0) = 40,000 ש"ח.
   * גריטה הפרשית נטו: 106,000 - 40,000 = **66,000 ש"ח**.
5. **חישוב NPV הפרשי ב-10%:**
   * $PV = \frac{177,333}{1.1^1} + \frac{177,333}{1.1^2} + \frac{177,333 + 66,000}{1.1^3}$
   * $PV = 161,212 + 146,556 + 182,820 = 490,588$ ש"ח.
   * $NPV = 490,588 - 590,000 = \mathbf{-99,412 \text{ ש\"ח}}$.
   * **החלטה:** לדחות את פרויקט ההחלפה (NPV הפרשי שלילי, ההשקעה משמידה ערך).

| תזרים הפרשי (בשקלים) | שנה 0 | שנה 1 | שנה 2 | שנה 3 |
| :--- | :---: | :---: | :---: | :---: |
| השקעה נטו במכונות | -590,000 | - | - | - |
| תזרים תפעולי הפרשי (OCF) | - | 177,333 | 177,333 | 177,333 |
| גריטה הפרשית נטו | - | - | - | 66,000 |
| **תזרים FCF חופשי הפרשי** | **-590,000** | **177,333** | **177,333** | **243,333** |
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Created Markdown file at: {output_path}")

def generate_docx(output_path):
    try:
        import docx
    except ImportError:
        print("docx package not installed. Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    # Define RTL styling helpers
    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_run_rtl(paragraph, text, bold=False, italic=False, font_size=11, color=None):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
        if color:
            run.font.color.rgb = color
        
        # Set RTL for run
        rPr = run._r.get_or_add_rPr()
        rBidi = OxmlElement('w:rtl')
        rBidi.set(qn('w:val'), '1')
        rPr.append(rBidi)
        return run

    # Document Title
    p = doc.add_paragraph()
    set_rtl(p)
    add_run_rtl(p, "בדיקת כדאיות השקעה: מדריך תיאורטי ותרגילים מעשיים", bold=True, font_size=18, color=RGBColor(30, 58, 138))
    
    p_sub = doc.add_paragraph()
    set_rtl(p_sub)
    add_run_rtl(p_sub, "חומר עזר ותירגול לסטודנטים במימון. המדריך מפרט את הכללים לקביעת תזרימי מזומנים רלוונטיים ומציג חמישה תרגילי כיתה פתורים, כולל חישובי NPV.", italic=True, font_size=11, color=RGBColor(100, 116, 139))

    # Goal
    p_goal = doc.add_paragraph()
    set_rtl(p_goal)
    add_run_rtl(p_goal, "עיקרון העל: עיקרון התזרים ההפרשי (Incremental Cash Flows)", bold=True, font_size=13, color=RGBColor(37, 99, 235))
    p_goal_text = doc.add_paragraph()
    set_rtl(p_goal_text)
    add_run_rtl(p_goal_text, "אנו כוללים בניתוח אך ורק את תזרימי המזומנים שישתנו בחברה כתוצאה ישירה מקבלת הפרויקט (תזרימים שינבעו מהפרויקט בהשוואה למצב בלעדיו).", font_size=11)

    # Section 1
    p_sec1 = doc.add_paragraph()
    set_rtl(p_sec1)
    add_run_rtl(p_sec1, "חלק 1: כללי אצבע - מה נכנס ומה לא נכנס לתחשיב כדאיות?", bold=True, font_size=14, color=RGBColor(30, 58, 138))

    p_in = doc.add_paragraph()
    set_rtl(p_in)
    add_run_rtl(p_in, "✓ מה נכנס לתחשיב? (תזרימים רלוונטיים)", bold=True, font_size=12, color=RGBColor(21, 128, 61))
    
    points_in = [
        ("עלויות השקעה ראשוניות (CapEx): ", "מחיר רכישת המכונה, הובלה, התקנה, הרצה והדרכת עובדים. כל אלו מהווים את בסיס הפחת של הנכס."),
        ("שינוי בהון חוזר נקי (NWC): ", "הגידול הנדרש במלאי או באשראי לקוחות (בניכוי אשראי ספקים) לצורך תמיכה בפרויקט. הון חוזר זה מוחזר במלואו כתזרים חיובי בשנה האחרונה."),
        ("עלויות אלטרנטיביות (Opportunity Costs): ", "שווי שוק של נכס קיים בבעלות החברה שישמש לפרויקט (לדוגמה, קרקע שניתן היה למכור או להשכיר). מחושב כתזרים שלילי נטו ממס בשנה 0."),
        ("קניבליזציה והשפעות לוואי: ", "ירידה צפויה במכירות של מוצרים קיימים של החברה עקב כניסת המוצר החדש."),
        ("מגן מס מפחת: ", "הפחת אינו תזרים מזומנים, אך הוא מקטין את חבות המס. מגן המס שווה ל-פחת שנתי × שיעור המס."),
        ("תזרים גריטה נטו ממס (Salvage Value): ", "המזומנים שיתקבלו ממכירת הנכס בסוף הפרויקט, בניכוי מס על רווח ההון (או בתוספת מגן מס במקרה של הפסד הון).")
    ]
    for title, desc in points_in:
        p_pt = doc.add_paragraph(style='List Bullet')
        set_rtl(p_pt)
        add_run_rtl(p_pt, title, bold=True)
        add_run_rtl(p_pt, desc)

    p_out = doc.add_paragraph()
    set_rtl(p_out)
    add_run_rtl(p_out, "✗ מה לא נכנס לתחשיב? (תזרימים לא רלוונטיים)", bold=True, font_size=12, color=RGBColor(185, 28, 28))

    points_out = [
        ("עלויות שקועות (Sunk Costs): ", "עלויות שכבר שולמו בעבר או שהחברה מחויבת להן בכל מקרה (למשל: סקר שוק, מחקר ופיתוח שכבר בוצע)."),
        ("הוצאות מימון (Financing Costs): ", "תשלומי ריבית, החזרי קרן הלוואה, או חלוקת דיבידנדים. (הסבר: עלויות אלו משוקללות כבר בתוך שיעור ההיוון/מחיר ההון WACC. הכללתן בתזרים תגרום לספירה כפולה)."),
        ("הקצאת עלויות תקורה קיימות: ", "העמסת הוצאות קבועות של החברה (כגון שכירות משרדי הנהלה, משכורות מנכ\"ל קיים) שאינן משתנות בעקבות הפרויקט."),
        ("הוצאות פחת כשלעצמן: ", "הפחת אינו יציאת מזומן פיזית. אנו משתמשים בו רק לצורך חישוב המס, ואז מוסיפים אותו בחזרה לרווח הנקי כדי לקבל את תזרים המזומנים התפעולי.")
    ]
    for title, desc in points_out:
        p_pt = doc.add_paragraph(style='List Bullet')
        set_rtl(p_pt)
        add_run_rtl(p_pt, title, bold=True)
        add_run_rtl(p_pt, desc)

    # Section 2 - Exercises
    p_sec2 = doc.add_paragraph()
    set_rtl(p_sec2)
    add_run_rtl(p_sec2, "חלק 2: תרגילי כיתה פתורים", bold=True, font_size=14, color=RGBColor(30, 58, 138))

    # Helper function to add exercises to docx
    def add_exercise_docx(title, description, solution_steps, table_data):
        p_title = doc.add_paragraph()
        set_rtl(p_title)
        add_run_rtl(p_title, title, bold=True, font_size=12, color=RGBColor(37, 99, 235))
        
        p_desc = doc.add_paragraph()
        set_rtl(p_desc)
        add_run_rtl(p_desc, description)
        
        p_sol_title = doc.add_paragraph()
        set_rtl(p_sol_title)
        add_run_rtl(p_sol_title, "פתרון:", bold=True, font_size=11, color=RGBColor(71, 85, 105))
        
        p_sol = doc.add_paragraph()
        set_rtl(p_sol)
        add_run_rtl(p_sol, solution_steps)
        
        # Add table
        t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        t.style = 'Light Shading Accent 1'
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                cell = t.cell(r_idx, c_idx)
                cell.text = val
                p_cell = cell.paragraphs[0]
                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                pPr = p_cell._p.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)
                
        doc.add_paragraph() # Spacing

    # Add Exercise 1
    ex1_title = "תרגיל 1: פרויקט הרחבה (הרמה הבסיסית)"
    ex1_desc = (
        "חברת \"סלטי השף\" שוקלת לרכוש קו ייצור חדש בעלות של 500,000 ש\"ח. הובלה והתקנה יעלו עוד 50,000 ש\"ח.\n"
        "• אורך הפרויקט: 4 שנים.\n"
        "• פחת: המכונה מופחתת בשיטת הקו הישר לאורך 5 שנים לערך ספרים אפס (קצב פחת של 20% מהעלות המקורית בשנה).\n"
        "• הון חוזר: בשנה 0 נדרש להגדיל את המלאי ב-40,000 ש\"ח. הון חוזר זה יוחזר במלואו בסוף שנה 4.\n"
        "• תוצאות תפעוליות: הכנסות שנתיות נוספות של 300,000 ש\"ח, והוצאות תפעוליות שנתיות (ללא פחת) של 120,000 ש\"ח.\n"
        "• סקר שוק שנערך בחודש שעבר בעלות של 30,000 ש\"ח הראה שיש ביקוש למוצר.\n"
        "• בסוף שנה 4 המכונה תימכר לצד ג' תמורת 150,000 ש\"ח.\n"
        "• שיעור מס חברות ומס רווח הון: 25%."
    )
    ex1_sol = (
        "1. בסיס הפחת של המכונה: 500,000 + 50,000 = 550,000 ש\"ח.\n"
        "2. הפחת השנתי: 550,000 × 20% = 110,000 ש\"ח לשנה.\n"
        "3. עלות שקועה: סקר השוק בסך 30,000 ש\"ח אינו רלוונטי ואינו נכלל בתזרים.\n"
        "4. ערך פנקסי בסוף שנה 4: 550,000 - (110,000 × 4) = 110,000 ש\"ח.\n"
        "5. תזרים גריטה נטו: 150,000 - 25% × (150,000 - 110,000) = 140,000 ש\"ח."
    )
    table_data_1 = [
        ["סעיף (בשקלים)", "שנה 0", "שנה 1", "שנה 2", "שנה 3", "שנה 4"],
        ["השקעה במכונה והתקנה", "-550,000", "-", "-", "-", "-"],
        ["שינוי בהון חוזר נקי", "-40,000", "-", "-", "-", "40,000"],
        ["הכנסות תפעוליות", "-", "300,000", "300,000", "300,000", "300,000"],
        ["הוצאות תפעוליות", "-", "-120,000", "-120,000", "-120,000", "-120,000"],
        ["פחת (לצרכי מס)", "-", "-110,000", "-110,000", "-110,000", "-110,000"],
        ["רווח תפעולי (EBIT)", "-", "70,000", "70,000", "70,000", "70,000"],
        ["מס חברות (25%)", "-", "-17,500", "-17,500", "-17,500", "-17,500"],
        ["רווח נקי", "-", "52,500", "52,500", "52,500", "52,500"],
        ["הוספת פחת בחזרה", "-", "110,000", "110,000", "110,000", "110,000"],
        ["תזרים תפעולי (OCF)", "-", "162,500", "162,500", "162,500", "162,500"],
        ["תזרים גריטה נטו", "-", "-", "-", "-", "140,000"],
        ["תזרים חופשי (FCF)", "-590,000", "162,500", "162,500", "162,500", "342,500"]
    ]
    add_exercise_docx(ex1_title, ex1_desc, ex1_sol, table_data_1)

    # Add Exercise 2
    ex2_title = "תרגיל 2: פרויקט החלפת מכונה (הרמה הבינונית)"
    ex2_desc = (
        "• המכונה הישנה: נרכשה לפני שנתיים ב-200,000 ש\"ח. מופחתת בקו ישר ל-5 שנים (40,000 ש\"ח פחת לשנה). כיום ניתנת למכירה ב-90,000 ש\"ח. אם תישאר, ערך גריטה שלה בעוד 3 שנים יהיה אפס.\n"
        "• המכונה החדשה: מחיר רכישה 300,000 ש\"ח. תופחת לאורך 3 שנים (חיי הפרויקט) לערך ספרים אפס (100,000 ש\"ח פחת לשנה). בסוף שנה 3 תימכר ב-50,000 ש\"ח.\n"
        "• שינוי תפעולי: המכונה החדשה תחסוך לחברה 80,000 ש\"ח בשנה בהוצאות התפעול.\n"
        "• שיעור מס חברות ומס רווח הון: 25%."
    )
    ex2_sol = (
        "1. השקעה נטו בשנה 0: -300,000 + 90,000 (מכירת הישנה) + 7,500 (מגן מס מהפסד הון על מכירת הישנה, שכן ערך ספרים שלה הוא 120,000 ש\"ח) = 202,500- ש\"ח.\n"
        "2. פחת הפרשי: 100,000 (חדשה) - 40,000 (ישנה) = 60,000 ש\"ח לשנה.\n"
        "3. תזרים תפעולי הפרשי (OCF): (80,000 - 60,000) × 75% + 60,000 = 75,000 ש\"ח.\n"
        "4. גריטה נטו בשנה 3: 50,000 - 25% × (50,000 - 0) = 37,500 ש\"ח."
    )
    table_data_2 = [
        ["תזרים הפרשי (בשקלים)", "שנה 0", "שנה 1", "שנה 2", "שנה 3"],
        ["רכישת מכונה חדשה", "-300,000", "-", "-", "-"],
        ["מכירת מכונה ישנה (כולל מס)", "97,500", "-", "-", "-"],
        ["חיסכון תפעולי שנתי", "-", "80,000", "80,000", "80,000"],
        ["מס על החיסכון והפחת", "-", "-5,000", "-5,000", "-5,000"],
        ["תזרים תפעולי הפרשי (OCF)", "-", "75,000", "75,000", "75,000"],
        ["גריטה נטו", "-", "-", "-", "37,500"],
        ["תזרים חופשי הפרשי (FCF)", "-202,500", "75,000", "75,000", "112,500"]
    ]
    add_exercise_docx(ex2_title, ex2_desc, ex2_sol, table_data_2)

    # Add Exercise 3
    ex3_title = "תרגיל 3: עלות אלטרנטיבית וקניבליזציה (הרמה המתקדמת)"
    ex3_desc = (
        "• השקעה ראשונית: 800,000 ש\"ח בשנה 0. מופחת בקו ישר ל-4 שנים לערך ספרים אפס.\n"
        "• עלות אלטרנטיבית: אובדן שכירות של 50,000 ש\"ח בשנה (לפני מס) בשנים 0, 1, 2 ו-3.\n"
        "• קניבליזציה: ירידה בתזרים המזומנים של מוצרים קיימים בסך 30,000 ש\"ח בשנה (שנים 1-4).\n"
        "• תוצאות תפעוליות של החדש: הכנסות 400,000 ש\"ח, הוצאות תפעוליות 150,000 ש\"ח.\n"
        "• שיעור מס: 20%."
    )
    ex3_sol = (
        "1. עלות אלטרנטיבית נטו: 50,000 × (1 - 0.20) = 40,000- ש\"ח בשנים 0-3.\n"
        "2. קניבליזציה נטו: 30,000- ש\"ח בשנים 1-4.\n"
        "3. OCF חדש: (250,000 - 200,000) × 80% + 200,000 = 240,000 ש\"ח."
    )
    table_data_3 = [
        ["סעיף (בשקלים)", "שנה 0", "שנה 1", "שנה 2", "שנה 3", "שנה 4"],
        ["השקעה בציוד ומפעל", "-800,000", "-", "-", "-", "-"],
        ["עלות אלטרנטיבית (שכירות נטו)", "-40,000", "-40,000", "-40,000", "-40,000", "-"],
        ["תזרים תפעולי מקונסולה חדשה", "-", "240,000", "240,000", "240,000", "240,000"],
        ["קניבליזציה", "-", "-30,000", "-30,000", "-30,000", "-30,000"],
        ["תזרים חופשי (FCF)", "-840,000", "170,000", "170,000", "170,000", "210,000"]
    ]
    add_exercise_docx(ex3_title, ex3_desc, ex3_sol, table_data_3)

    # Add Exercise 4
    ex4_title = "תרגיל 4: השקת מוצר חדש עם חישוב NPV (חדש!)"
    ex4_desc = (
        "حברת \"גיימינג פרו\" שוקלת לרכוש ציוד לייצור בקרים חדשים. עלות הציוד היא 1,200,000 ש\"ח, ועלויות התקנה והרצה הן 100,000 ש\"ח.\n"
        "• אורך הפרויקט: 5 שנים.\n"
        "• פחת: שיטת קו ישר לאורך 5 שנים לערך ספרים 0 (פחת שנתי של 260,000 ש\"ח).\n"
        "• הון חוזר נקי (NWC): נדרשת השקעה ראשונית של 80,000 ש\"ח בשנה 0, אשר תוחזר במלואה בסוף שנה 5.\n"
        "• נתונים תפעוליות: הכנסות של 600,000 ש\"ח בשנה, והוצאות תפעוליות (ללא פחת) של 200,000 ש\"ח בשנה.\n"
        "• קניבליזציה: ירידה בתזרים המזומנים של מוצרים קיימים בגובה 20,000 ש\"ח בשנה (לאחר מס).\n"
        "• גריטה: הציוד יימכר בסוף שנה 5 ב-200,000 ש\"ח.\n"
        "• שיעור המס: 23%.\n"
        "• מחיר ההון של הפרויקט (שיעור ההיוון): 12%."
    )
    ex4_sol = (
        "1. שנה 0 (השקעה נטו): -1,200,000 (ציוד) - 100,000 (התקנה) - 80,000 (NWC) = 1,380,000- ש\"ח.\n"
        "2. פחת שנתי: 1,300,000 / 5 = 260,000 ש\"ח.\n"
        "3. תזרים תפעולי (שנים 1-5): EBITDA = 400,000 ש\"ח. EBIT = 140,000 ש\"ח. מס = 32,200 ש\"ח. רווח נקי = 107,800 ש\"ח. OCF מחושב = רווח נקי + פחת = 367,800 ש\"ח. OCF כולל קניבליזציה = 347,800 ש\"ח.\n"
        "4. גריטה נטו בשנה 5: 200,000 - 23% × (200,000 - 0) = 154,000 ש\"ח.\n"
        "5. תזרים מזומנים שנה 5 כולל: 347,800 (תפעולי) + 154,000 (גריטה) + 80,000 (החזר הון חוזר) = 581,800 ש\"ח.\n"
        "6. חישוב NPV ב-12%:\n"
        "PV = 347,800/(1.12^1) + 347,800/(1.12^2) + 347,800/(1.12^3) + 347,800/(1.12^4) + 581,800/(1.12^5)\n"
        "PV = 310,536 + 277,264 + 247,557 + 221,033 + 330,129 = 1,386,519 ש\"ח.\n"
        "NPV = 1,386,519 - 1,380,000 = +6,519 ש\"ח.\n"
        "החלטה: ה-NPV חיובי ולכן הפרויקט כדאי."
    )
    table_data_4 = [
        ["סעיף (בשקלים)", "שנה 0", "שנה 1", "שנה 2", "שנה 3", "שנה 4", "שנה 5"],
        ["ציוד והתקנה", "-1,300,000", "-", "-", "-", "-", "-"],
        ["השקעה בהון חוזר נקי", "-80,000", "-", "-", "-", "-", "80,000"],
        ["תזרים תפעולי (OCF) נטו", "-", "347,800", "347,800", "347,800", "347,800", "347,800"],
        ["גריטה נטו ממס", "-", "-", "-", "-", "-", "154,000"],
        ["תזרים FCF חופשי", "-1,380,000", "347,800", "347,800", "347,800", "347,800", "581,800"]
    ]
    add_exercise_docx(ex4_title, ex4_desc, ex4_sol, table_data_4)

    # Add Exercise 5
    ex5_title = "תרגיל 5: פרויקט החלפת ציוד עם חישוב NPV (חדש!)"
    ex5_desc = (
        "חברה שוקלת להחליף מכונה ישנה במכונה חדשה וחסכונית.\n"
        "• המכונה הישנה: נרכשה לפני 3 שנים ב-600,000 ש\"ח. היא מופחתת בקו ישר ל-6 שנים (100,000 ש\"ח לשנה). כיום ניתנת למכירה ב-250,000 ש\"ח. אם תישאר, ערך הגריטה שלה בעוד 3 שנים יהיה 50,000 ש\"ח.\n"
        "• המכונה החדשה: עלות רכישה והתקנה היא 850,000 ש\"ח בשנה 0. מופחתת בקו ישר לאורך 3 שנים לערך ספרים של 50,000 ש\"ח. בסוף שנה 3 היא תימכר ב-120,000 ש\"ח.\n"
        "• שינוי תפעולי: המכונה החדשה תחסוך להוצאות התפעול 180,000 ש\"ח בשנה (לפני מס).\n"
        "• שיעור מס: 20%.\n"
        "• מחיר ההון (שיעור ההיוון): 10%."
    )
    ex5_sol = (
        "1. השקעה נטו בשנה 0: -850,000 + 250,000 (מכירה) + 10,000 (מגן מס מהפסד הון על הישנה) = 590,000- ש\"ח.\n"
        "2. פחת הפרשי: פחת חדשה (266,667) - פחת ישנה (100,000) = 166,667 ש\"ח לשנה.\n"
        "3. תזרים תפעולי הפרשי (OCF): (180,000 - 166,667) × 80% + 166,667 = 177,333 ש\"ח לשנה.\n"
        "4. גריטה הפרשית נטו בשנה 3: גריטה נטו חדשה (106,000 ש\"ח) - גריטה נטו ישנה שנמנעה (40,000 ש\"ח) = 66,000 ש\"ח.\n"
        "5. חישוב NPV הפרשי ב-10%:\n"
        "PV = 177,333/(1.1^1) + 177,333/(1.1^2) + (177,333 + 66,000)/(1.1^3)\n"
        "PV = 161,212 + 146,556 + 182,820 = 490,588 ש\"ח.\n"
        "NPV = 490,588 - 590,000 = -99,412 ש\"ח.\n"
        "החלטה: ה-NPV ההפרשי שלילי. ההשקעה אינה כדאית לחברה ומומלץ להמשיך להשתמש במכונה הישנה."
    )
    table_data_2_diff = [
        ["תזרים הפרשי (בשקלים)", "שנה 0", "שנה 1", "שנה 2", "שנה 3"],
        ["השקעה נטו במכונות", "-590,000", "-", "-", "-"],
        ["תזרים תפעולי הפרשי (OCF)", "-", "177,333", "177,333", "177,333"],
        ["גריטה הפרשית נטו", "-", "-", "-", "66,000"],
        ["תזרים FCF חופשי הפרשי", "-590,000", "177,333", "177,333", "243,333"]
    ]
    add_exercise_docx(ex5_title, ex5_desc, ex5_sol, table_data_2_diff)

    doc.save(output_path)
    print(f"Created Word Document at: {output_path}")

if __name__ == "__main__":
    folder_path = "C:\\Users\\aviti\\.gemini\\antigravity-ide\\scratch\\finance-course-materials"
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    
    html_path = os.path.join(folder_path, "capital_budgeting_guide.html")
    md_path = os.path.join(folder_path, "capital_budgeting_guide.md")
    docx_path = os.path.join(folder_path, "capital_budgeting_guide.docx")
    
    create_html_file(html_path)
    create_markdown_file(md_path)
    generate_docx(docx_path)
    print("All file generations completed successfully!")
